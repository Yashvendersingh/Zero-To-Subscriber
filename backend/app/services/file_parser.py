import io
import logging
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        full_text = "\n".join(text_parts)
        if not full_text.strip():
            raise ValueError("PDF appears to be image-based or contains no extractable text")
        return full_text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Could not extract text from PDF: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts)
        if not full_text.strip():
            raise ValueError("DOCX appears to contain no extractable text")
        return full_text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Could not extract text from DOCX: {str(e)}")

def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception as e:
            raise ValueError(f"Could not decode text file: {str(e)}")

async def parse_resume_file(file_bytes: bytes, filename: str) -> str:
    """Parse a resume file and return its text content."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    elif ext == "txt":
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Please upload a PDF, DOCX, or TXT file.")
